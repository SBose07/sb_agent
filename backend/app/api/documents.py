from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from typing import List
import io
import re
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..models.document import Document, DocumentCreate, DocumentUpdate
from ..services.document_service import document_service


router = APIRouter(tags=["documents"])


@router.get("/", response_model=List[Document])
async def get_documents():
    """Get all documents."""
    return document_service.get_all()


@router.get("/{document_id}", response_model=Document)
async def get_document(document_id: str):
    """Get a specific document by ID."""
    document = document_service.get_by_id(document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return document


@router.post("/", response_model=Document)
async def create_document(data: DocumentCreate):
    """Create a new document."""
    return document_service.create(data)


@router.put("/{document_id}", response_model=Document)
async def update_document(document_id: str, data: DocumentUpdate):
    """Update an existing document."""
    document = document_service.update(document_id, data)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return document


@router.delete("/{document_id}")
async def delete_document(document_id: str):
    """Delete a document."""
    success = document_service.delete(document_id)
    if not success:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"status": "deleted", "id": document_id}


@router.post("/upload", response_model=Document)
async def upload_docx(file: UploadFile = File(...)):
    """
    Upload a .docx file and create a document from its content.
    The content is extracted as plain text with basic formatting preserved.
    """
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
    
    try:
        # Read the uploaded file
        contents = await file.read()
        docx_file = io.BytesIO(contents)
        
        # Parse the docx file
        docx = DocxDocument(docx_file)
        
        # Extract text with basic markdown formatting
        lines = []
        for paragraph in docx.paragraphs:
            text = paragraph.text.strip()
            if not text:
                lines.append("")
                continue
            
            # Convert heading styles to markdown
            style_name = paragraph.style.name.lower() if paragraph.style else ""
            if 'heading 1' in style_name:
                text = f"# {text}"
            elif 'heading 2' in style_name:
                text = f"## {text}"
            elif 'heading 3' in style_name:
                text = f"### {text}"
            elif 'heading 4' in style_name:
                text = f"#### {text}"
            elif 'title' in style_name:
                text = f"# {text}"
            
            lines.append(text)
        
        # Also extract tables
        for table in docx.tables:
            lines.append("")
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    # Add header separator
                    lines.append("|" + "|".join(["---" for _ in cells]) + "|")
            lines.append("")
        
        content = '\n'.join(lines)
        
        # Get filename without extension as title
        title = file.filename.rsplit('.', 1)[0]
        
        # Create document
        document = document_service.create(DocumentCreate(title=title, content=content))
        return document
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")


@router.get("/{document_id}/download")
async def download_docx(document_id: str):
    """
    Download a document as a .docx file.
    Markdown formatting is converted to Word styles.
    """
    document = document_service.get_by_id(document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    try:
        # Create a new Word document
        docx = DocxDocument()
        
        # Parse markdown content and add to document
        lines = document.content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # Check for table (starts with |)
            if stripped.startswith('|') and '|' in stripped[1:]:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                
                # Parse and create table
                create_word_table(docx, table_lines)
                continue
            
            # Check for code block
            if stripped.startswith('```'):
                code_lines = []
                lang = stripped[3:].strip()
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                i += 1  # Skip closing ```
                
                # Add code as formatted paragraph
                code_text = '\n'.join(code_lines)
                p = docx.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                p.paragraph_format.left_indent = Inches(0.5)
                continue
            
            if not stripped:
                # Empty paragraph
                docx.add_paragraph("")
                i += 1
                continue
            
            # Handle markdown headings
            if stripped.startswith('#### '):
                docx.add_heading(stripped[5:], level=4)
            elif stripped.startswith('### '):
                docx.add_heading(stripped[4:], level=3)
            elif stripped.startswith('## '):
                docx.add_heading(stripped[3:], level=2)
            elif stripped.startswith('# '):
                docx.add_heading(stripped[2:], level=1)
            elif stripped.startswith('> '):
                # Blockquote as italic with indent
                p = docx.add_paragraph()
                run = p.add_run(stripped[2:])
                run.italic = True
                p.paragraph_format.left_indent = Inches(0.5)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                # Bullet list
                docx.add_paragraph(stripped[2:], style='List Bullet')
            elif len(stripped) > 2 and stripped[0].isdigit() and (stripped[1] == '.' or (stripped[1].isdigit() and stripped[2] == '.')):
                # Numbered list
                text = stripped.split('.', 1)[1].strip() if '.' in stripped else stripped
                docx.add_paragraph(text, style='List Number')
            elif stripped == '---' or stripped == '***':
                # Horizontal rule - add empty paragraph with bottom border
                docx.add_paragraph("_" * 50)
            else:
                # Regular paragraph - handle inline formatting
                p = docx.add_paragraph()
                add_formatted_text(p, stripped)
            
            i += 1
        
        # Save to bytes buffer
        buffer = io.BytesIO()
        docx.save(buffer)
        buffer.seek(0)
        
        # Return as downloadable file
        filename = f"{document.title}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating file: {str(e)}")


def create_word_table(docx, table_lines: List[str]):
    """Convert markdown table lines to a Word table."""
    if len(table_lines) < 2:
        return
    
    # Parse table rows
    rows = []
    for line in table_lines:
        # Skip separator lines (|---|---|)
        if re.match(r'^\|[\s\-:]+\|$', line.replace('|', '|').replace('-', '-')):
            continue
        if '---' in line and line.count('|') > 2:
            continue
            
        # Parse cells
        cells = [cell.strip() for cell in line.split('|')]
        # Remove empty first/last elements from leading/trailing |
        cells = [c for c in cells if c or cells.index(c) not in [0, len(cells)-1]]
        cells = [c for c in cells if c]  # Remove empty strings
        
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    # Determine number of columns
    num_cols = max(len(row) for row in rows)
    
    # Create table
    table = docx.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Fill in cells
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                # Handle inline formatting in cells
                cell.text = ""
                p = cell.paragraphs[0]
                add_formatted_text(p, cell_text)
                
                # Make header row bold
                if i == 0:
                    for run in p.runs:
                        run.bold = True
    
    # Add empty paragraph after table
    docx.add_paragraph("")


def add_formatted_text(paragraph, text: str):
    """Add text to paragraph with basic markdown formatting (bold, italic, code)."""
    # Pattern for **bold**, *italic*, and `code`
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)
